import docx
import sys

def add_content_to_docx(file_path, heading, content_file):
    """
    Adds a heading and content from a file to a .docx file.

    Args:
        file_path (str): The path to the .docx file.
        heading (str): The heading to add.
        content_file (str): The path to the file containing the content.
    """
    try:
        with open(content_file, 'r') as f:
            content = f.read()
    except Exception as e:
        print(f"Error reading content file: {e}")
        sys.exit(1)

    try:
        document = docx.Document(file_path)
    except Exception as e:
        print(f"Error opening document: {e}")
        sys.exit(1)

    document.add_heading(heading, level=1)
    document.add_paragraph(content)

    try:
        document.save(file_path)
        print(f"Successfully added content to {file_path}")
    except Exception as e:
        print(f"Error saving document: {e}")
        sys.exit(1)

if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("Usage: python edit_docx.py <file_path> <heading> <content_file>")
        sys.exit(1)

    file_path = sys.argv[1]
    heading = sys.argv[2]
    content_file = sys.argv[3]
    add_content_to_docx(file_path, heading, content_file)